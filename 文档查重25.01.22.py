

"""
基于OOP模式开发，分为验证部分和主程序两部分
11.16 共有匹配逻辑修改为：全局完整出现：如果某个匹配内容在每个txt文件中都至少完整出现了一次，那么它应该被保存。
局部完整出现与全局子字符串出现：如果某个匹配内容至少在一个txt文件中完整出现，并且是其他每个txt文件中匹配内容的子字符串，那么它也应该被记录。
"""



#12.26删除灰色导入
import multiprocessing
from concurrent.futures import ProcessPoolExecutor, ThreadPoolExecutor
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
import smtplib
import random
import json
from subprocess import check_output
from tkinter import simpledialog
from multiprocessing import Manager
from threading import Thread
import numpy as np
from tkinter import filedialog, messagebox
from queue import Queue, Empty
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from shutil import rmtree
import hashlib
from pathlib import Path
import imagehash
from PIL import Image
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import tkinter as tk
from tkinter import Toplevel, Label, Checkbutton
from tkinter import ttk
import string
from concurrent.futures import ThreadPoolExecutor, as_completed
import logging
import psutil
from collections import defaultdict
import gc
import re
import requests
import os
import shutil
import sys
import subprocess
from tensorflow.keras.applications.resnet50 import ResNet50, preprocess_input
from tensorflow.keras.preprocessing import image as keras_image
from sklearn.metrics.pairwise import cosine_similarity
import tensorflow as tf





# 配置日志
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

# 减少 TensorFlow 日志输出
os.environ['TF_CPP_MIN_LOG_LEVEL'] = '2'






class EmailVerifier:
    def __init__(self):
        self.sender_email = "416391746@qq.com"
        self.recipient_email = "416391746@qq.com"
        self.smtp_server = "smtp.qq.com"
        self.smtp_port = 465
        self.smtp_username = "416391746@qq.com"
        self.smtp_password = "lkaunakpwgxfcahe"
        self.verify_status_file = "验证通过记录（勿删）.json"

    def generate_code(self):
        return random.randint(100000, 999999)

    def send_verification_code(self, email, code):
        try:
            msg = MIMEMultipart()
            msg['From'] = self.sender_email
            msg['To'] = email
            msg['Subject'] = "验证码"
            body = f"您的验证码为: {code}"
            msg.attach(MIMEText(body, 'plain'))

            server = smtplib.SMTP_SSL(self.smtp_server, self.smtp_port)
            server.login(self.smtp_username, self.smtp_password)
            server.send_message(msg)
            server.quit()
            print(f"温馨提示：已发送验证码，请索要验证码！")
        except smtplib.SMTPAuthenticationError:
            print("认证错误，请检查用户名和密码是否正确。")
        except smtplib.SMTPServerDisconnected as e:
            print(f"连接意外关闭: {e}")
        except Exception as e:
            print(f"发送邮件时发生错误: {e}")

    def get_hard_disk_serial(self):
        try:
            output = check_output('wmic diskdrive get SerialNumber', shell=True, text=True)
            lines = output.split('\n')
            for line in lines:
                if line.strip().startswith('SerialNumber'):
                    return line.split()[-1].strip()
            return None
        except Exception as e:
            #print(f"获取硬盘序列号失败,无法正常验证: {e}")
            print(f"计算机信息更新失败,无法正常验证: {e}")
            return None

    def generate_hash_from_serial(self, serial):
        if serial:
            return hashlib.sha256(serial.encode()).hexdigest()
        return None

    def record_verification_status(self, verified=True, hard_disk_hash=None):
        try:
            with open(self.verify_status_file, 'w') as file:
                #data = {'verified': verified, 'hard_disk_hash': hard_disk_hash}
                data = {'verified': verified, 'secret': hard_disk_hash}
                json.dump(data, file)
        except Exception as e:
            print(f"写入验证状态文件时发生错误: {e}")

    def ask_for_code_and_verify(self, root):
        code = self.generate_code()
        self.send_verification_code(self.recipient_email, code)
        entered_code = simpledialog.askstring("验证码验证", "请输入您得到的验证码:")
        if entered_code == str(code):
            hard_disk_serial = self.get_hard_disk_serial()
            if hard_disk_serial:
                hard_disk_hash = self.generate_hash_from_serial(hard_disk_serial)
                self.record_verification_status(True, hard_disk_hash)
                #messagebox.showinfo("验证成功", "验证码和硬盘序列号验证成功！")
                messagebox.showinfo("验证成功", "信息验证成功！")
                self.send_email_notification(self.recipient_email, code, True)

                # 关闭当前窗口
                root.destroy()

                # 创建新的根窗口并启动主程序
                new_root = tk.Tk()
                new_root.title("文档查重工具")
                same_word_app = DocumentApp(new_root)
                new_root.mainloop()

            else:
                #messagebox.showerror("验证失败", "无法获取硬盘序列号，验证失败。")
                messagebox.showerror("验证失败", "无法更新计算机信息，验证失败。")
        else:
            messagebox.showerror("验证失败", "验证码错误，请重新验证。")
            self.send_email_notification(self.recipient_email, code, False)
            # 检查是否已经存在“重新发送验证码”按钮，如果存在则移除
            for widget in root.winfo_children():
                if isinstance(widget, tk.Button) and widget['text'] == "重新发送验证码":
                    widget.destroy()
                    break

            # 创建一个新的“重新发送验证码”按钮
            request_new_code_button = tk.Button(root, text="重新发送验证码",
                                                command=lambda: self.ask_for_code_and_verify(root))
            request_new_code_button.pack(pady=10)


    def send_email_notification(self, email, code, success):
        subject = "验证码验证完成" if success else "验证码验证失败"
        message = f"您的验证码为: {code}已验证成功" if success else f"未输入正确的验证码'{code}'导致验证失败"

        msg = MIMEMultipart()
        msg['From'] = self.sender_email
        msg['To'] = email
        msg['Subject'] = subject
        msg.attach(MIMEText(message, 'plain'))

        try:
            server = smtplib.SMTP_SSL(self.smtp_server, self.smtp_port)
            server.login(self.smtp_username, self.smtp_password)
            server.sendmail(self.sender_email, email, msg.as_string())
            server.quit()
            print("邮件发送成功。")
        except Exception as e:
            print(f"发送邮件时发生错误: {e}")

    def check_verification_status(self):
        try:
            if os.path.exists(self.verify_status_file):
                with open(self.verify_status_file, 'r') as file:
                    try:
                        status = json.load(file)
                        verified = status.get('verified', False)
                        #hard_disk_hash = status.get('hard_disk_hash')
                        hard_disk_hash = status.get('secret')
                        if verified and hard_disk_hash:
                            current_hard_disk_serial = self.get_hard_disk_serial()
                            if current_hard_disk_serial:
                                current_hard_disk_hash = self.generate_hash_from_serial(current_hard_disk_serial)
                                return current_hard_disk_hash == hard_disk_hash
                    except json.JSONDecodeError:
                        print("不能通过复制验证文件，请通过正规途径进行验证。")
                        return False
            return False
        except Exception as e:
            print(f"读取验证状态文件时发生错误: {e}")
            return False

#主程序

class DocumentComparer:
    def __init__(self, min_chars, tolerance, hamming_threshold, similarity, batch_size, file_paths, output_path, filter_chars=None):
        self.min_chars = min_chars
        self.tolerance = tolerance
        self.hamming_threshold = hamming_threshold
        self.similarity = similarity
        self.batch_size = batch_size
        self.file_paths = file_paths
        self.output_path = output_path
        self.completed_comparisons = 0  # 添加为类实例变量
        self.interval_tree_root = None


        # 创建一个Manager实例用于多进程共享对象
        self.manager = Manager()
        self.lock = self.manager.Lock()
        self.documents = self.manager.dict()  # 使用 Manager 的 dict
        self.paragraph_positions = self.manager.dict()  # 使用 Manager 的 dict
        self.matches_all = self.manager.dict()  # 使用 Manager 的 dict
        self.common_matches = self.manager.list()  # 使用 Manager 的 list
        self.filtered_matches = self.manager.list() #初始化时设置为 manager 的 list
        self.filtered_matches_event = self.manager.Event()  # 使用 Manager 的 Event


        self.common_matches_created_event = self.manager.Event()  # 使用 Manager 的 Event
        self.lock = self.manager.Lock()  # 引入锁机制，用于保护共享资源

        self.matches_created_event = self.manager.Event()  # 文字比较完成事件
        self.images_compared_event = self.manager.Event()  # 图片比较完成事件
        self.tables_compared_event = self.manager.Event()  # 表格比较完成事件


        # 初始化图片和表格比对相关属性
        self.common_images = self.manager.list()
        self.common_tables = self.manager.dict()
        self.image_model = ResNet50(weights='imagenet', include_top=False, pooling='avg')

        self.filter_chars = filter_chars  # 存储过滤字符集








    def load_documents(self):
        unique_file_paths = list(dict.fromkeys(self.file_paths))  # 去重
        self.file_paths = unique_file_paths
        for fp in self.file_paths:
            try:
                doc = Document(fp)
                paragraphs = [p.text for p in doc.paragraphs]
                positions = []
                current_pos = 0
                for paragraph in paragraphs:
                    end_pos = current_pos + len(paragraph)
                    positions.append((current_pos, end_pos))
                    current_pos = end_pos + 1
                self.documents[fp] = paragraphs
                self.paragraph_positions[fp] = positions
                logging.info(f"Loaded document {fp}")
            except Exception as e:
                logging.error(f"Failed to load document {fp}: {e}")
                raise



    #以下为图片比对
    @staticmethod
    def _get_image_features(img_path, model):
        try:
            img = keras_image.load_img(img_path, target_size=(224, 224))
            logging.info(f"Extracting features from image: {img_path}")
            x = keras_image.img_to_array(img)
            x = np.expand_dims(x, axis=0)
            x = preprocess_input(x)
            features = model.predict(x)

            return features.flatten()
        except Exception as e:
            logging.error(f"Error extracting features from image {img_path}: {e}")
            return None

    @staticmethod
    def are_images_identical_by_bytes(img1_path, img2_path):
        """通过比较文件的字节内容判断两张图片是否完全相同"""
        img1_path = Path(img1_path)
        img2_path = Path(img2_path)

        if not img1_path.is_file() or not img2_path.is_file():
            logging.error(f"One of the files does not exist: {img1_path} or {img2_path}")
            return False

        # 检查文件大小是否相同
        if img1_path.stat().st_size != img2_path.stat().st_size:
            logging.info(f"Files have different sizes: {img1_path} and {img2_path}")
            return False

        # 分块读取文件内容以节省内存并逐块比较
        with img1_path.open('rb') as f1, img2_path.open('rb') as f2:
            while True:
                chunk1 = f1.read(8192)
                chunk2 = f2.read(8192)
                if not chunk1 or not chunk2:
                    break  # 文件读取完毕
                if chunk1 != chunk2:
                    logging.info(f"Files differ at a chunk: {img1_path} and {img2_path}")
                    return False  # 如果任意一块不相等，则文件不同

        logging.info(f"Files are identical: {img1_path} and {img2_path}")
        return True



    @staticmethod
    def _compare_images_phash(img1_path, img2_path, hamming_threshold=None):#5为比较折衷的数字，越小，表示两张图像越相似，可以挑出来有细微差别的图片，包括压缩、裁剪、旋转
        try:
            #logging.info(f"Comparing images using pHash: {img1_path} and {img2_path}")
            logging.info(f"Comparing images using pHash: {img1_path} and {img2_path}")

            # 打开并计算两张图片的pHash
            with Image.open(img1_path) as img1, Image.open(img2_path) as img2:
                hash1 = imagehash.phash(img1)
                hash2 = imagehash.phash(img2)

            # 计算汉明距离
            hamming_distance = hash1 - hash2
            #logging.info(f"Hamming distance between {img1_path} and {img2_path}: {hamming_distance}")
            logging.info(f"Hamming distance between {img1_path} and {img2_path}: {hamming_distance}")

            # 使用提供的hamming_threshold或默认值5
            threshold = hamming_threshold if hamming_threshold is not None else 5

            #print(f"这是使用的hamming_threshold={threshold}")

            # 根据汉明距离判断相似度
            return hamming_distance <= threshold
            #print(f"这是使用的hamming_threshold={threshold}")
        except Exception as e:
            logging.error(f"Error comparing images {img1_path} and {img2_path} using Hamming: {e}")
            return False



    @staticmethod
    def _compare_images(img1_path, img2_path, model, similarity_threshold=None):
        try:
            #logging.info(f"Comparing images: {img1_path} and {img2_path}")
            # 首先尝试通过字节级比较确定图片是否完全相同
            if DocumentComparer.are_images_identical_by_bytes(img1_path, img2_path):
                #logging.info(f"Images {img1_path} and {img2_path} are identical according to byte comparison.")
                logging.info(f"Images {img1_path} and {img2_path} are identical according to 100%-OK comparison.")
                return True



            # 使用pHash方法进行比较，并从外部传入hamming_threshold参数
            hamming_threshold = getattr(DocumentComparer, 'hamming_threshold', 5)  # 获取类级别的默认值或使用5作为默认
            if DocumentComparer._compare_images_phash(img1_path, img2_path, hamming_threshold=hamming_threshold):
                logging.info(f"Images {img1_path} and {img2_path} are similar according to hamming comparison.")
                #print(f"这是使用的hamming_threshold 第二处={hamming_threshold}")
                return True

            # 如果提供了模型并且需要进一步使用深度学习模型进行特征比对
            if model is not None:
                # 如果需要，继续使用深度学习模型进行特征比对
                img1_features = DocumentComparer._get_image_features(img1_path, model)
                img2_features = DocumentComparer._get_image_features(img2_path, model)
                # 使用余弦相似度来衡量两个特征向量的相似性
                similarity = cosine_similarity([img1_features], [img2_features])[0][0]
                #logging.info(f"Similarity between {img1_path} and {img2_path}: {similarity}")
                logging.info(f"Similarity between {img1_path} and {img2_path}: {similarity}")
                threshold = similarity_threshold if similarity_threshold is not None else 0.9
                #print(f"这是使用的similarity_threshold={threshold}")

                return similarity >= threshold  # 使用传递的相似度阈值或默认值
            # 如果没有提供模型，则仅依靠前两种方法
            return False
        except Exception as e:
            logging.error(f"Error comparing images {img1_path} and {img2_path}: {e}")
            return False

    @staticmethod
    def compare_image_similarity_and_find_common(images1, images2, doc1, doc2, common_images_shared, lock, model, similarity_threshold=None):
        for img1_info in images1:
            for img2_info in images2:
                if DocumentComparer._compare_images(img1_info[0], img2_info[0], model, similarity_threshold=similarity_threshold):
                    with lock:
                        common_images_shared.append(((doc1, img1_info[0]), (doc2, img2_info[0])))
                        logging.info(f"Found matching images: {doc1} image {img1_info[0]}, {doc2} image {img2_info[0]}")

    def extract_images_from_doc(self, doc_path, temp_dir=None):
        document = Document(doc_path)
        images_info = []  # 存储图片路径及其对应的文档路径


        try:
            # 创建基于文档名称的子目录来存储该文档中的所有图片
            doc_name = os.path.basename(doc_path).rsplit('.', 1)[0]  # 去掉扩展名
            #doc_dir = os.path.join(self.output_path, doc_name)
            doc_dir = os.path.join(self.output_path, f'{doc_name}-表格与图片')
            #images_subdir = os.path.join(doc_dir, 'images')
            images_subdir = os.path.join(doc_dir, '图片')
            os.makedirs(images_subdir, exist_ok=True)

            for rel in document.part.rels.values():
                if "image" in rel.target_ref:
                    img_partname = rel._target.partname.lstrip('/')
                    img_blob = rel._target.blob
                    img_filename = os.path.basename(img_partname)

                    # 使用唯一标识符作为文件名的一部分，避免冲突
                    unique_img_name = f"{img_filename}"
                    img_save_path = os.path.join(images_subdir, unique_img_name)

                    with open(img_save_path, 'wb') as img_file:
                        img_file.write(img_blob)

                    images_info.append((img_save_path, doc_path))  # 记录图片路径和文档路径
                    logging.info(f"Extracted and saved image to: {img_save_path}")

        except Exception as e:
            logging.error(f"Failed to extract images from {doc_path}: {e}", exc_info=True)

        return images_info

    def copy_image(self, src, dst):
        try:
            shutil.copyfile(src, dst)
            logging.info(f"Copied image from {src} to {dst}")
        except Exception as e:
            logging.error(f"Failed to copy image from {src} to {dst}: {e}", exc_info=True)

    def compare_and_save_common_images(self, progress_queue, similarity_threshold, batch_size):
        try:
            all_images = {doc_path: self.extract_images_from_doc(doc_path) for doc_path in self.file_paths}

            # 提前创建输出目录
            images_output_dir = Path(self.output_path) / '重复图片'
            os.makedirs(images_output_dir, exist_ok=True)

            image_pairs = [
                ((doc1, images1), (doc2, images2))
                for i, (doc1, images1) in enumerate(all_images.items())
                for j, (doc2, images2) in enumerate(all_images.items())
                if i < j
            ]

            total_tasks = len(image_pairs)
            completed_tasks = 0

            with ProcessPoolExecutor(max_workers=max((os.cpu_count() or 1) // 2, 1)) as process_pool:
                # 分批提交任务到进程池
                for start in range(0, total_tasks, batch_size // 10):
                    batch_pairs = image_pairs[start:start + batch_size // 10]
                    # 使用进程池并行执行特征提取和比较任务
                    futures = [
                        process_pool.submit(
                            self.compare_image_similarity_and_find_common,
                            images1, images2, doc1, doc2, self.common_images, self.lock, self.image_model,
                            similarity_threshold
                        )
                        for (doc1, images1), (doc2, images2) in batch_pairs
                    ]

                    for future in as_completed(futures):
                        future.result()
                        completed_tasks += 1
                        progress = int(95 * completed_tasks / total_tasks)
                        progress_queue.put(progress)



            pair_counter = defaultdict(int)  # 用来记录每对文档之间的相似图片数量

            # 使用线程池处理图片的写入操作
            with ThreadPoolExecutor(max_workers=max((os.cpu_count() or 1) * 2, 1)) as thread_pool:
                if self.common_images:
                    futures = []
                    for (doc1, img_path1), (doc2, img_path2) in list(self.common_images):
                        doc1_name = Path(doc1).stem
                        doc2_name = Path(doc2).stem
                        pair_folder_name = f'{doc1_name}-{doc2_name}-重复图片'
                        pair_folder_path = images_output_dir / pair_folder_name
                        os.makedirs(pair_folder_path, exist_ok=True)

                        # 为每对文档之间的相似图片分配一个唯一的标识符
                        docs_tuple = tuple(sorted([doc1, doc2]))
                        pair_id = f"pair{pair_counter[docs_tuple]:03d}"
                        pair_counter[docs_tuple] += 1

                        # 构建新的文件名，包含配对编号和原始图片名称，并保留原始图片扩展名
                        image_filename1 = Path(img_path1)
                        image_filename2 = Path(img_path2)

                        # 使用 pair_folder_path 作为基础路径来保存共同图片
                        output_image_path1 = pair_folder_path / f'{pair_id}：{doc1_name}({image_filename1.stem}){image_filename1.suffix}'
                        output_image_path2 = pair_folder_path / f'{pair_id}：{doc2_name}({image_filename2.stem}){image_filename2.suffix}'

                        # 提交写入任务到线程池
                        futures.append(thread_pool.submit(self.copy_image, img_path1, output_image_path1))
                        futures.append(thread_pool.submit(self.copy_image, img_path2, output_image_path2))
                    # 等待所有拷贝操作完成
                    for future in as_completed(futures):
                        future.result()

            # 清理临时文件夹
            temp_dir = Path(self.output_path) / 'temp_images'
            if temp_dir.exists():
                rmtree(temp_dir)
                logging.info(f"Removed temporary directory: {temp_dir}")

            with self.lock:
                progress_queue.put(100)  # Update progress to 100%

            # 清理模型会话，确保资源释放
            tf.keras.backend.clear_session()

            logging.info("Image comparison completed.")
        except Exception as e:
            logging.error("Error during image comparison and saving common images.", exc_info=True)
            with self.lock:
                progress_queue.put(85)  # Update progress to 85% even if there is an error

    #以下为表格比对
    @staticmethod
    def _is_table_empty(data):
        """检查表格是否为空"""
        return all(not any(cell.strip() for cell in row) for row in data)

    @staticmethod
    def _convert_table_to_2d_list(table):
        """将表格对象转换为二维列表，并去除所有空白字符，但保留所有列"""

        def remove_all_whitespace(text):
            """去除字符串中的所有空白字符"""
            if not isinstance(text, str):
                text = str(text)
            return re.sub(r'\s+', '', text)  # 使用正则表达式去除所有空白字符

        if isinstance(table, list) and all(isinstance(row, list) for row in table):
            return [[remove_all_whitespace(cell) for cell in row] for row in table if any(remove_all_whitespace(cell) for cell in row)]
        elif hasattr(table, 'rows'):
            return [[remove_all_whitespace(cell.text) for cell in row.cells] for row in table.rows if
                    any(remove_all_whitespace(cell.text) for cell in row.cells)]
        else:
            raise ValueError("Unsupported table format")

    @staticmethod
    def _compare_tables(table1, table2, tolerance=None):
        try:

            # # 记录原始表格内容
            # original_data1 = [
            #     [cell.text.strip() if hasattr(cell, 'text') else str(cell).strip()
            #      for cell in row.cells]  # 确保访问的是单元格，而不是行对象
            #     for row in table1.rows if hasattr(table1, 'rows')
            # ]
            # original_data2 = [
            #     [cell.text.strip() if hasattr(cell, 'text') else str(cell).strip()
            #      for cell in row.cells]  # 确保访问的是单元格，而不是行对象
            #     for row in table2.rows if hasattr(table2, 'rows')
            # ]
            #
            # logging.info("Original data1 content: %s", original_data1)
            # logging.info("Original data2 content: %s", original_data2)

            # 将表格内容转换为二维列表，并去除空白字符
            data1 = DocumentComparer._convert_table_to_2d_list(table1)
            data2 = DocumentComparer._convert_table_to_2d_list(table2)

            # 如果任一表格为空，则直接返回False
            if DocumentComparer._is_table_empty(data1) or DocumentComparer._is_table_empty(data2):
                logging.info("One of the tables is empty.")
                return False

            # 检查表格是否有足够的有效数据（至少两行）
            if len(data1) <= 1 or len(data2) <= 1:
                logging.info("One of the tables has only one row after filtering; comparison not meaningful.")
                return False

            # 确保data1是较小的表格
            if len(data1) * (len(data1[0]) if data1 else 0) > len(data2) * (len(data2[0]) if data2 else 0):
                data1, data2 = data2, data1  # 如果data2更小，则交换

            # 使用提供的容差或默认值0.9
            threshold = tolerance if tolerance is not None else 0.9

            # 辅助函数：计算两个子表格之间的相似度
            def calculate_similarity(sub_data1, sub_data2):
                # 计算所有单元格的总数
                total_cells = len(sub_data1) * (len(sub_data1[0]) if sub_data1 else 0)
                # 计算匹配单元格数量
                matching_cells = sum(
                    int(cell1.strip() == cell2.strip())
                    for row1, row2 in zip(sub_data1, sub_data2)
                    for cell1, cell2 in zip(row1, row2)
                )

                # 返回相似度
                return matching_cells / total_cells if total_cells > 0 else 1.0

            # 滑动窗口遍历较大表格的所有可能位置
            rows_diff = max(0, len(data2) - len(data1) + 1)
            cols_diff = max(0, (len(data2[0]) - len(data1[0]) + 1) if data1 and data2 else 0)

            window_data = []  # 重用这个列表来保存窗口数据
            for i in range(rows_diff):
                for j in range(cols_diff):
                    # 提取当前窗口的数据
                    window_data[:]= [row[j:j + len(data1[0])] for row in data2[i:i + len(data1)]]

                    # 记录 window_data 内容
                    #logging.info(f"Extracted window_data at ({i},{j}): {window_data}")

                    #logging.info(f"data1 length: {len(data1)}, data2 length: {len(data2)}")
                    #logging.info(f"data1 content: {data1}")
                    #logging.info(f"data2 content: {data2}")

                    # 确保提取的窗口数据和data1有相同的尺寸
                    if all(len(w_row) == len(data1[0]) for w_row in window_data):
                        similarity = calculate_similarity(data1, window_data)

                        #logging.info(f"data1: {data1}")
                        #logging.info(f"window_data: {window_data}")

                        #logging.info(f"Sliding window at ({i},{j}) with tolerance similarity: {similarity}. Target: {threshold}")
                        logging.info(f"Sliding ({i},{j}) with tolerance similarity: {similarity}. Target: {threshold}")
                        if similarity >= threshold:
                            del window_data[:]  # 清空窗口数据
                            return True  # 一旦找到满足条件的匹配，立即返回True并终止循环
                    # 手动清理不再需要的对象
                    del window_data[:]


            # # 特别处理只在外部增加行或列的情况
            # if len(data1) < len(data2) or any(len(row1) < len(row2) for row1, row2 in zip(data1, data2)):
            #     # 如果data2是在data1基础上增加了行或列，检查重叠部分的相似度
            #     overlap_rows = min(len(data1), len(data2))
            #     overlap_cols = min(len(data1[0]), len(data2[0])) if data1 and data2 else 0
            #     overlap_data1 = [row[:overlap_cols] for row in data1[:overlap_rows]]
            #     overlap_data2 = [row[:overlap_cols] for row in data2[:overlap_rows]]
            #
            #     similarity = calculate_similarity(overlap_data1, overlap_data2)
            #     logging.info(f"Checking overlap region with similarity: {similarity}. Target: {threshold}")
            #     if similarity >= threshold:
            #         return True


            return False  # 如果没有找到满足条件的匹配，则返回False
        except Exception as e:
            logging.error(f"Error comparing tables: {e}")
            return False

    @staticmethod
    def compare_tables(doc1_path, doc2_path,common_tables_shared, lock,tolerance):
        doc1 = Document(doc1_path)
        doc2 = Document(doc2_path)
        tables1 = list(doc1.tables)
        tables2 = list(doc2.tables)
        common_pairs = []
        for i, table1 in enumerate(tables1):
            for j, table2 in enumerate(tables2):
                if DocumentComparer._compare_tables(table1, table2,tolerance=tolerance):
                    common_pairs.append(((doc1_path, i), (doc2_path, j)))
        with lock:
            if (doc1_path, doc2_path) not in common_tables_shared:
                common_tables_shared[(doc1_path, doc2_path)] = common_pairs
            else:
                common_tables_shared[(doc1_path, doc2_path)].extend(common_pairs)

    @staticmethod
    def copy_table_with_style(src_table, dest_doc):
        # 使用 len(src_table.columns) 来获取列数
        new_table = dest_doc.add_table(rows=0, cols=len(src_table.columns))

        # 深拷贝每一行和单元格的内容及样式
        for row in src_table.rows:
            new_row = new_table.add_row()
            for idx, cell in enumerate(row.cells):
                new_cell = new_row.cells[idx]
                for paragraph in cell.paragraphs:
                    new_paragraph = new_cell.paragraphs[0] if len(new_cell.paragraphs) > 0 else new_cell.add_paragraph()
                    for run in paragraph.runs:
                        new_run = new_paragraph.add_run(run.text)
                        # 复制字体样式
                        new_run.bold = run.bold
                        new_run.italic = run.italic
                        new_run.underline = run.underline
                        if run.font.color:
                            new_run.font.color.rgb = run.font.color.rgb
                        if run.font.size:
                            new_run.font.size = run.font.size
                        # 其他样式属性也可以类似地复制

                # 如果有其他非段落内容（如图片），也应一并处理

        return new_table

    @staticmethod
    def add_page_break(doc):
        """
        在Word文档中添加一个分页符。
        """
        page_break = OxmlElement('w:p')
        page_break.append(OxmlElement('w:r'))
        page_break_run = OxmlElement('w:br')
        page_break_run.set(qn('w:type'), 'page')
        page_break.append(page_break_run)
        doc.element.body.append(page_break)






    def extract_tables_from_doc(self, doc_path):
        document = Document(doc_path)
        tables_info = []  # 存储表格路径及其对应的文档路径
        all_tables_save_path = None  # 初始化为空，确保有默认值

        try:
            # 创建基于文档名称的子目录来存储该文档中的所有表格
            doc_name = os.path.basename(doc_path).rsplit('.', 1)[0]  # 去掉扩展名
            #doc_dir = os.path.join(self.output_path, doc_name)
            doc_dir = os.path.join(self.output_path,  f'{doc_name}-表格与图片')
            #tables_subdir = os.path.join(doc_dir, 'tables')
            tables_subdir = os.path.join(doc_dir, '表格')
            os.makedirs(tables_subdir, exist_ok=True)

            # 创建一个新的 DOCX 文件来保存所有表格
            all_tables_doc = Document()

            def add_title_single_table(table_idx):
                """ 添加表格编号作为标题 """
                # 创建段落并添加标题文本
                para = all_tables_doc.add_paragraph()
                title_text = f"{doc_name}(表格 {table_idx})"
                run = para.add_run(title_text)

                # 设置标题样式（例如，加粗、左对齐等）
                run.bold = True
                run.font.size = Pt(12)
                para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                #logging.info(f"Added title for table {table_idx}: {title_text}")

            for idx, table in enumerate(document.tables, start=1):

                # 添加表格编号作为标题
                add_title_single_table(idx)
                # 使用深拷贝方法保留原始格式
                self.copy_table_with_style(table, all_tables_doc)

                # 添加一个空段落作为间隔
                all_tables_doc.add_paragraph("")

                # 在添加下一个表格前插入分页符（除了最后一个表格）
                if idx < len(document.tables):
                    DocumentComparer.add_page_break(all_tables_doc)


                # 直接使用 idx 记录日志
                #logging.info(f"Extracted and added table {idx} from document {doc_path}")


            # 保存包含所有表格的 DOCX 文件
            all_tables_save_path = os.path.join(tables_subdir, f"{doc_name}_所有表格.docx")
            all_tables_doc.save(all_tables_save_path)
            logging.info(f"Saved all extracted tables to: {all_tables_save_path}")

        except Exception as e:
            logging.error(f"Failed to extract tables from {doc_path}: {e}", exc_info=True)

        return all_tables_save_path if all_tables_save_path else None


    def compare_and_save_common_tables(self, progress_queue,tolerance,batch_size):
        try:
            all_tables_indices = {}  # 存储文档路径及其对应的表格索引列表
            extracted_tables_info = {}  # 存储已提取表格的信息

            # 提取所有表格索引
            for doc_path in self.file_paths:
                document = Document(doc_path)
                all_tables_indices[doc_path] = list(range(len(document.tables)))
                # 调用 extract_tables_from_doc 方法来单独保存表格
                tables_info = self.extract_tables_from_doc(doc_path)
                extracted_tables_info[doc_path] = tables_info

            # 创建用于保存共同表格的目录
            tables_output_dir = Path(self.output_path) / '重复表格'
            tables_output_dir.mkdir(parents=True, exist_ok=True)

            table_pairs = [
                ((doc1, doc2))
                for doc1 in self.file_paths
                for doc2 in self.file_paths
                if doc1 < doc2  # 确保只进行一次双向比较
            ]

            total_tasks = len(table_pairs)
            completed_tasks = 0


            with multiprocessing.Pool(processes=max((os.cpu_count() or 1) // 2, 1)) as pool:
                for start in range(0, total_tasks, batch_size//5):
                    batch_pairs = table_pairs[start:start + batch_size//5]
                    pool.starmap(
                            DocumentComparer.compare_tables,
                            [(doc1, doc2, self.common_tables, self.lock, tolerance)
                             for (doc1, doc2) in batch_pairs]  # 注意这里修正了参数解包方式
                        )
                    completed_tasks += len(batch_pairs)
                    progress = int(75 * completed_tasks / total_tasks)
                    progress_queue.put(progress)

                    # 每个批次处理完后手动触发垃圾回收
                    gc.collect()

            # 对于每一对文档，创建一个单独的Word文档来保存它们的共同表格
            with ThreadPoolExecutor(max_workers=max((os.cpu_count() or 1) * 2, 1)) as executor:
                futures = []
                for (doc1, doc2), common_tables_list in self.common_tables.items():
                    if common_tables_list:
                        futures.append(executor.submit(
                            DocumentComparer.save_common_tables,
                            doc1, doc2, common_tables_list, all_tables_indices, tables_output_dir
                        ))

                for future in as_completed(futures):
                    future.result()  # 确保异常被捕获

            # 在所有任务完成后再次触发垃圾回收，确保所有临时资源被清理
            gc.collect()

            with self.lock:
                progress_queue.put(85)  # 更新进度到100%
            logging.info("Table comparison completed.")
        except Exception as e:
            logging.error("Error during table comparison and saving common tables.", exc_info=True)
            # 即使发生错误，也尝试进行垃圾回收
            gc.collect()
        finally:
            with self.lock:
                progress_queue.put(70)  # 更新进度到70%

    @staticmethod
    def save_common_tables(doc1, doc2, common_tables_list, all_tables_indices, tables_output_dir):
        common_table_doc = Document()
        pair_counter = 0

        def add_title_for_table(doc_name, table_index, pair_counter):
            para = common_table_doc.add_paragraph()
            run = para.add_run(
                f"Pair{str(pair_counter).zfill(3)}:{doc_name}(表格{table_index + 1})")
            run.bold = True
            run.font.size = Pt(12)
            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            #logging.info(f"Added title for {doc_name} table {table_index + 1}")

        doc1_name = Path(doc1).stem
        doc2_name = Path(doc2).stem
        for idx, ((_, table_index1), (_, table_index2)) in enumerate(common_tables_list):
            pair_counter += 1


            # 添加第一个表格的标题
            add_title_for_table(doc1_name, table_index1, pair_counter)

            # 紧跟标题添加第一个表格
            doc1_tables = Document(doc1).tables
            common_table1 = DocumentComparer.copy_table_with_style(doc1_tables[table_index1], common_table_doc)
            #logging.info(f"Copied table from {doc1} index {table_index1}")

            # 添加第二个表格的标题
            add_title_for_table(doc2_name, table_index2, pair_counter)

            # 紧跟标题添加第二个表格
            doc2_tables = Document(doc2).tables
            common_table2 = DocumentComparer.copy_table_with_style(doc2_tables[table_index2], common_table_doc)
            #logging.info(f"Copied table from {doc2} index {table_index2}")

            # 添加一个空段落作为间隔
            common_table_doc.add_paragraph("")

            # 在这对表格之后插入分页符（除了最后一对表格）
            if idx < len(common_tables_list) - 1:
                DocumentComparer.add_page_break(common_table_doc)
                #logging.info(f"Added page break after pair {pair_counter}")

        # 保存这个文档对的共同表格到文件
        output_file = tables_output_dir / f'{doc1_name}-{doc2_name}-重复表格.docx'
        common_table_doc.save(output_file)
        logging.info(f"Saved common tables to: {output_file}")







    #以下是文字比较部分，为了多进程加这里进行修改

    @staticmethod
    def process_paragraphs(k, l, doc1, doc2, positions1, positions2, min_chars, filter_chars):
        local_matches = []
        paragraph_matches = DocumentComparer.sliding_window_compare(doc1[k], doc2[l], min_chars, filter_chars=filter_chars)
        for match in paragraph_matches:
            global_start1 = positions1[k][0] + match[0]
            global_end1 = positions1[k][0] + match[1]
            global_start2 = positions2[l][0] + match[2]
            global_end2 = positions2[l][0] + match[3]
            local_matches.append((global_start1, global_end1, global_start2, global_end2, match[4]))
        return local_matches

    @staticmethod
    def compare_two_docs_static(documents, file_paths, paragraph_positions, i, j, batch_size=None,min_chars=30, filter_chars='', output_path='', matches_all_dict=None, common_matches_list=None, lock=None):
        # 这个函数现在是类的一个方法，可以被pickle
        #logging.info(f"#### TEXT COMPARISON STARTED FOR DOCUMENTS {i} AND {j} ####")
        matches = []
        doc1 = documents[file_paths[i]]
        doc2 = documents[file_paths[j]]
        positions1 = paragraph_positions[file_paths[i]]
        positions2 = paragraph_positions[file_paths[j]]


        # 分批提交任务
        total_tasks = len(doc1) * len(doc2)
        tasks = [(k, l) for k in range(len(doc1)) for l in range(len(doc2))]
        batch_size = batch_size if batch_size is not None else 100


        for start in range(0, total_tasks, batch_size):
            batch_tasks = tasks[start:start + batch_size]
            for k, l in batch_tasks:
                try:
                    batch_matches = DocumentComparer.process_paragraphs(k, l, doc1, doc2, positions1, positions2,
                                                                        min_chars, filter_chars)
                    matches.extend(batch_matches)
                except Exception as e:
                    logging.error(f"Error comparing paragraphs {k},{l}: {e}")

            logging.info(
                f"Batch completed at task {start // (batch_size or 100)}. Current memory usage: {psutil.Process().memory_info().rss / 1024 / 1024:.2f} MB")

            # 手动触发垃圾回收并记录内存使用情况
            #del batch_tasks  # 删除对批次任务列表的引用
            #gc.collect()  # 手动触发垃圾回收


        file1_name = os.path.basename(file_paths[i])
        file2_name = os.path.basename(file_paths[j])
        DocumentComparer.save_matches(matches, file1_name, file2_name, output_path)
        if matches and matches_all_dict is not None and common_matches_list is not None and lock is not None:
            DocumentComparer.update_matches_all(matches, matches_all_dict, i, j, lock)




        return matches, i, j


    def compare_documents(self, progress_queue,**kwargs):
        # 使用kwargs中的参数覆盖实例属性
        self.min_chars = kwargs.get('min_chars', self.min_chars)
        self.tolerance = kwargs.get('tolerance', self.tolerance)
        self.hamming_threshold = kwargs.get('hamming_threshold', self.hamming_threshold)
        self.similarity_threshold = kwargs.get('similarity_threshold', self.similarity)
        self.filter_chars = kwargs.get('filter_chars', self.filter_chars)
        self.batch_size = kwargs.get('batch_size', self.batch_size)

        # 记录启动比较时的参数值
        logging.info(f"Starting comparison with parameters: "
                     f"min_chars={self.min_chars}, "
                     f"tolerance={self.tolerance}, "
                     f"hamming_threshold={self.hamming_threshold}, "
                     f"similarity_threshold={self.similarity_threshold}, "  # 使用 similarity_threshold
                     f"filter_chars={self.filter_chars},"
                     f"batch_size={self.batch_size}")



        try:
            self.load_documents()
            logging.info(f"File paths after loading: {self.file_paths}")  # 添加调试日志

            all_results = []

            total_comparisons = len(self.file_paths) * (len(self.file_paths) - 1) // 2
            completed_comparisons = 0

            comparer_data = {
                'documents': self.documents,
                'file_paths': self.file_paths,
                'paragraph_positions': self.paragraph_positions,
                'min_chars': self.min_chars,
                'filter_chars': self.filter_chars,
                'output_path': self.output_path,  # 确保 output_path 是类的一个属性
                'matches_all_dict': self.matches_all,
                'common_matches_list': self.common_matches,
                'lock': self.lock
            }

            with multiprocessing.Pool(processes=max((os.cpu_count() or 1) // 2, 1)) as pool:#取cpu的一半
                results = [
                    pool.apply_async(DocumentComparer.compare_two_docs_static, args=(
                        comparer_data['documents'], comparer_data['file_paths'], comparer_data['paragraph_positions'],
                        i, j, self.batch_size, comparer_data['min_chars'], comparer_data['filter_chars'], comparer_data['output_path'],
                        comparer_data['matches_all_dict'],
                        comparer_data['common_matches_list'],
                        comparer_data['lock']
                    ))
                    for i in range(len(self.file_paths))
                    for j in range(i + 1, len(self.file_paths))
                ]

                for result in results:
                    try:
                        matches, i, j= result.get()
                        if isinstance(matches, list):
                            all_results.append((matches, i, j))
                            completed_comparisons += 1
                            progress = int(50 * completed_comparisons / total_comparisons)
                            progress_queue.put(progress)
                        else:
                            # logging.error("Unexpected return type from compare_two_docs")
                            logging.error("Unexpected return.")
                    except Exception as e:
                        logging.error(f"Error during comparison: {e}")


            # 确保所有文本比较任务完成后才继续
            with self.lock:
                self.matches_created_event.set()  # 设置文本比较完成事件
                progress_queue.put(70)  # 确保进度条更新到70%
            logging.info("Text comparison completed.")

            # 调用保存共同匹配内容的方法
            self.save_common_matches()

            # 调用保存过滤后匹配内容的方法
            self.save_filtered_matches()

            self.filtered_matches_event.wait()

            # 表格比较任务
            self.compare_and_save_common_tables(progress_queue, self.tolerance, self.batch_size)
            with self.lock:
                #self.tables_compared_event.set()  # 标记表格比较已完成  多余
                progress_queue.put(85)  # 确保进度条更新到100%
            logging.info("Table comparison completed.")


            # 图片比较任务
            self.compare_and_save_common_images(progress_queue, self.similarity_threshold, self.batch_size)
            with self.lock:
                #self.images_compared_event.set()  # 标记图片比较已完成  多余
                progress_queue.put(100)  # 更新进度条到85%
            logging.info("Image comparison completed.")



        except Exception as e:
            logging.error(f"An error occurred in compare_documents: {e}")



    @staticmethod
    def _filter_text(text: str, filter_chars: set = None) -> str:
        """过滤文本中指定的字符"""
        if filter_chars is None:
            return text  # 如果没有指定过滤字符，则不进行过滤
        return ''.join([char for char in text if char not in filter_chars])

    @staticmethod
    def sliding_window_compare(text1: str, text2: str, min_chars: int,filter_chars: set = None) -> list:
        if not isinstance(text1, str) or not isinstance(text2, str) or not (
                isinstance(min_chars, int) and min_chars > 0):
            raise ValueError("Invalid input types or values.")

        # 对输入文本进行过滤
        # 对输入文本进行过滤，并记录原始和过滤后的文本长度
        original_len1, original_len2 = len(text1), len(text2)

        if original_len1 < min_chars or original_len2 < min_chars:
            #logging.info("text lengths are less than min_chars, no comparison needed.")
            return []  # 直接返回空列表，无需进一步比较

        filtered_text1 = DocumentComparer._filter_text(text1, filter_chars)
        filtered_text2 = DocumentComparer._filter_text(text2, filter_chars)
        #logging.info(f"Original lengths: {original_len1}, {original_len2}")
        #logging.info(f"Filtered lengths: {len(filtered_text1)}, {len(filtered_text2)}")
        #print(f"这是使用的filter_chars={filter_chars}")

        # 12.27 新增检查过滤后的文本长度是否满足最小字符数要求
        if len(filtered_text1) < min_chars or len(filtered_text2) < min_chars:
            #logging.info("Filtered text lengths are less than min_chars, no comparison needed.")
            return []  # 直接返回空列表，无需进一步比较

        matches = []
        len_filtered_text1, len_filtered_text2 = len(filtered_text1), len(filtered_text2)
        #logging.info(f"Starting comparison between text1 and text2 with min_chars={min_chars}...")


        def is_submatch(existing_match, new_match):
            """检查新匹配是否被现有匹配完全包含"""
            (start1_existing, end1_existing, start2_existing, end2_existing, _) = existing_match
            (start1_new, end1_new, start2_new, end2_new, _) = new_match
            return (start1_existing <= start1_new < end1_existing and
                    start2_existing <= start2_new < end2_existing)

        #print(f"这是使用的min_chars={min_chars}")
        i = 0
        while i <= len_filtered_text1 - min_chars:
            for j in range(len_filtered_text2 - min_chars + 1):
                if filtered_text1[i:i + min_chars] == filtered_text2[j:j + min_chars]:
                    end1, end2 = i + min_chars, j + min_chars
                    while end1 < len_filtered_text1 and end2 < len_filtered_text2 and filtered_text1[end1] == filtered_text2[end2]:
                        end1 += 1
                        end2 += 1
                    match_text = filtered_text1[i:end1]

                    # 尝试合并当前匹配到已有的匹配中
                    merged = False
                    for idx, (start1, end1_existing, start2, end2_existing, _) in enumerate(matches):
                        # 检查是否可以合并
                        if (start1 <= i < end1_existing) and (start2 <= j < end2_existing) or \
                                (i <= start1 < end1) and (j <= start2 < end2) or \
                                (start1 < i < end1_existing) and (start2 < j < end2_existing) or \
                                (i < start1 < end1) and (j < start2 < end2):
                            new_start1 = min(i, start1)
                            new_end1 = max(end1, end1_existing)
                            new_start2 = min(j, start2)
                            new_end2 = max(end2, end2_existing)
                            matches[idx] = (new_start1, new_end1, new_start2, new_end2, filtered_text1[new_start1:new_end1])
                            merged = True
                            break

                    if not merged:
                        # 在添加新的匹配之前，检查它是否已经被记录
                        new_match = (i, end1, j, end2, match_text)
                        if not any(is_submatch(existing, new_match) for existing in matches):
                            matches.append(new_match)
                            #logging.info(f"Found match: {match_text} from {i} to {end1} in text1 and from {j} to {end2} in text2")
                            logging.info(f"Found match: {match_text}")


            # 在完成对text1当前位置的所有可能匹配后，移动到下一个位置
            i += 1


        logging.info(f"Comparison complete, found {len(matches)} matches.")

        return matches

    @staticmethod
    def update_matches_all(matches, matches_all_dict, doc_index1, doc_index2, lock):
        with lock:  # 确保同一时间只有一个进程更新 shared data
            for match in matches:
                match_text = match[4]
                start1 = match[0]
                start2 = match[2]
                if match_text not in matches_all_dict:
                    matches_all_dict[match_text] = {doc_index1: (start1, match[1]), doc_index2: (start2, match[3])}
                    #logging.info(f"Added new match_text: {match_text} to matches_all")
                    logging.info(f"Added new match_text: {match_text}")

                else:
                    matches_all_dict[match_text][doc_index1] = (start1, match[1])
                    matches_all_dict[match_text][doc_index2] = (start2, match[3])
                    #logging.info(f"Updated existing match_text: {match_text} in matches_all")
                    #logging.info(f"Updated existing match_text")

                    # 检查是否有同一位置的更长匹配内容
                    for existing_match in list(matches_all_dict.keys()):
                        if match_text in existing_match and match_text != existing_match:
                            existing_start_pos, existing_end_pos = matches_all_dict[existing_match].get(doc_index1,
                                                                                                        (None, None))
                            if existing_start_pos == start1 and existing_end_pos == match[1]:
                                del matches_all_dict[match_text]
                                #logging.info(f"Removed match_text: {match_text} from matches_all due to longer match at the same position")
                                logging.info(f"Removed match_text: {match_text}")

                                break

    @staticmethod
    def save_matches(matches, file1_name, file2_name, output_path):
        # 去掉 .docx 后缀
        file1_name = os.path.splitext(file1_name)[0]  # 修改点：去掉 .docx 后缀
        file2_name = os.path.splitext(file2_name)[0]  # 修改点：去掉 .docx 后缀

        output_file = os.path.join(output_path, f"{file1_name}-{file2_name}[重复内容].txt")
        logging.info(f"Attempting to write to {output_file}")

        try:
            os.makedirs(os.path.dirname(output_file), exist_ok=True)

            if not matches:
                with open(output_file, 'w', encoding='utf-8') as f:
                    pass  # 创建空文件
                logging.info(f"Created an empty matches file: {output_file}")
                return
            # 按匹配内容的长度降序排序，然后按起始位置升序排序
            matches = sorted(matches, key=lambda x: (len(x[4]), x[0], x[2]), reverse=True)
            logging.info(f"Sorted matches by length and position")
            # 去重：确保同一位置的最长匹配内容优先
            unique_matches = []
            for match in matches:
                if not any(
                        match[0] >= m[0] and match[1] <= m[1] and match[2] >= m[2] and match[3] <= m[3] for m in
                        unique_matches):
                    unique_matches.append(match)
                    #logging.info(f"Added unique match: {match[4]}")
                    #logging.info(f"Added unique match: {match[4]}")


            # 用于记录已经写入的匹配文本
            written_texts = set()

            with open(output_file, 'w', encoding='utf-8') as f:
                for match in unique_matches:
                    match_text = match[4]
                    start1, end1, start2, end2 = match[:4]
                    if match_text not in written_texts:
                        f.write(f"{match_text}\n")
                        # f.write(f"Start position in {file1_name}: {start1}\n")
                        # f.write(f"End position in {file1_name}: {end1}\n")
                        # f.write(f"Start position in {file2_name}: {start2}\n")
                        # f.write(f"End position in {file2_name}: {end2}\n")
                        # f.write("-" * 40 + "\n")
                        written_texts.add(match_text)


            logging.info(f"Saved matches to {output_file}")


        except Exception as e:
            logging.error(f"Failed to save matches to {output_file}: {e}")




    def save_common_matches(self):
        self.matches_created_event.wait()  # 等待事件被设置
        # 获取所有生成的匹配文件
        match_files = [
            os.path.join(self.output_path,f"{os.path.splitext(os.path.basename(file1))[0]}-{os.path.splitext(os.path.basename(file2))[0]}[重复内容].txt")
            for i, file1 in enumerate(self.file_paths)
            for j, file2 in enumerate(self.file_paths)
            if i < j
        ]

        # 初始化字典来存储每个匹配内容的出现情况
        match_presence = {}
        for match_file in match_files:
            try:
                if not os.path.exists(match_file):
                    continue
                with open(match_file, 'r', encoding='utf-8') as f:
                    for line in f:
                        line = line.strip()
                        if line:
                            if line not in match_presence:
                                match_presence[line] = {'full_matches': 0, 'substring_matches': 0, 'files': set()}
                            match_presence[line]['files'].add(match_file)
            except Exception as e:
                #logging.error(f"Failed to read match file {match_file}: {e}")
                logging.error(f"Failed to read match file {match_file}: {e}")


        # 检查每个匹配内容是否全局完整出现或局部完整出现且为其他文件的子字符串
        for match_file in match_files:
            if not os.path.exists(match_file):
                continue
            with open(match_file, 'r', encoding='utf-8') as f:
                content = f.read()
                for match, data in match_presence.items():
                    if match in content:
                        data['full_matches'] += 1
                    elif any(match in substring for substring in content.split()):
                        data['substring_matches'] += 1

        with self.lock:  # 确保同一时间只有一个进程更新 shared data
            matches_to_save = []
            for match, data in match_presence.items():
                if data['full_matches'] == len(match_files) or (
                        data['full_matches'] >= 1 and data['substring_matches'] == len(match_files) - data[
                    'full_matches']):
                    matches_to_save.append(match)
            # 按照匹配内容长度降序排序
            sorted_matches = sorted(matches_to_save, key=len, reverse=True)

            # 更新 common_matches 列表
            self.common_matches[:] = sorted_matches

        # 保存共同匹配内容到文件
        self.save_common_matches_to_file()
        self.common_matches_created_event.set()  # 设置事件

    def save_common_matches_to_file(self):
        output_file = os.path.join(self.output_path, "所有文档的重复内容.txt")
        try:
            with open(output_file, 'w', encoding='utf-8') as f:
                for match_text in self.common_matches:
                    f.write(f"{match_text}\n")
            # 记录日志
            if self.common_matches:
                logging.info(f"Saved common matches to {output_file}")
            else:
                logging.info(f"Created an empty common matches file: {output_file}")
        except Exception as e:
            logging.error(f"Failed to save common matches to {output_file}: {e}")



    def save_filtered_matches(self):
        self.common_matches_created_event.wait()  # 等待所有 save_common_matches 文件生成完成
        # 获取所有生成的匹配文件
        match_files = [
            os.path.join(self.output_path,f"{os.path.splitext(os.path.basename(file1))[0]}-{os.path.splitext(os.path.basename(file2))[0]}[重复内容].txt")
            for i, file1 in enumerate(self.file_paths)
            for j, file2 in enumerate(self.file_paths)
            if i < j
        ]

        # 遍历所有匹配文件，生成新的输出文件，剔除掉所有共同部分的内容
        for match_file in match_files:
            try:
                with open(match_file, 'r', encoding='utf-8') as f:
                    lines = f.readlines()

                filtered_lines = [line for line in lines if line.strip() not in self.common_matches]
                new_output_file = match_file.replace("[重复内容].txt", "[重复内容（去除所有文档的重复内容）].txt")
                with open(new_output_file, 'w', encoding='utf-8') as f:
                    for line in filtered_lines:
                        f.write(line)
                # 记录日志
                if filtered_lines:
                    logging.info(f"Saved filtered matches to {new_output_file}")
                else:
                    logging.info(f"Created an empty filtered matches file: {new_output_file}")
            except Exception as e:
                logging.error(f"Failed to process match file {match_file}: {e}")


        # 设置过滤匹配完成事件
        with self.lock:
            self.filtered_matches_event.set()  # 标记过滤匹配完成




    def start_comparison(self, progress_queue):
        self.progress_queue = progress_queue
        self.compare_documents()


class DocumentApp:
    def __init__(self, root):
        self.root = root
        self.root.title("文档查重工具")
        # 文件列表
        self.file_paths = []
        # 输出路径
        self.output_path = ""  # 初始化output_path属性为空字符串
        # 显示信息的文本框
        self.info_text = tk.Text(root, wrap='word', height=10)
        self.info_text.pack(pady=10)
        # 文件选择按钮
        self.file_button = tk.Button(root, text="选择文件", command=self.select_files)
        self.file_button.pack(pady=10)
        # 清空文件列表按钮
        self.clear_files_button = tk.Button(root, text="清空文件列表", command=self.clear_files)
        self.clear_files_button.pack(pady=10)
        # 添加设置按钮
        self.settings_button = tk.Button(root, text="个性化设置", command=self.open_settings)
        self.settings_button.pack(pady=10)

        # 初始化默认值
        self.min_chars = tk.StringVar(value="50")
        self.tolerance = tk.StringVar(value="0.9")
        self.hamming_threshold = tk.StringVar(value="5")
        self.similarity = tk.StringVar(value="0.9")
        self.filter_chars_var = tk.StringVar()  # 用于管理过滤字符的StringVa
        self.batch_size = tk.StringVar(value="100")  # 默认批处理大小
        self.filter_chars = set()  # 初始化过滤字符为空集合
        self.remove_punctuation = tk.BooleanVar(value=False)  # 是否移除所有标点符号


        # 输出路径选择按钮
        self.output_button = tk.Button(root, text="选择输出路径", command=self.select_output_path)
        self.output_button.pack(pady=10)
        # 开始按钮
        self.start_button = tk.Button(root, text="开始比较", command=self.start_comparison)
        self.start_button.pack(pady=10)
        # 进度条
        self.progress_var = tk.IntVar()
        self.progress_bar = ttk.Progressbar(root, variable=self.progress_var, maximum=100)
        self.progress_bar.pack(pady=10)

        self.qq_label = ttk.Label(root, text="QQ：1011489402", foreground="blue", cursor="hand2")
        self.qq_label.pack( pady= 10)

    def create_tooltip(self, widget, text):
        def enter(event):
            x = widget.winfo_rootx() + widget.winfo_width() + 5
            y = widget.winfo_rooty() + widget.winfo_height() // 2

            # Creates a toplevel window for the tooltip
            self.tooltip_window = Toplevel(widget)
            self.tooltip_window.wm_overrideredirect(True)  # Removes the border of the window
            self.tooltip_window.wm_geometry(f"+{x}+{y}")

            label = Label(self.tooltip_window, text=text, background="#ffffe0", relief="solid", borderwidth=1)
            label.pack()

        def leave(event):
            if hasattr(self, 'tooltip_window') and self.tooltip_window:
                self.tooltip_window.destroy()
                delattr(self, 'tooltip_window')

        widget.bind("<Enter>", enter)
        widget.bind("<Leave>", leave)

    def open_settings(self):
        settings_window = tk.Toplevel(self.root)
        settings_window.title("个性化设置")


        tooltips = {
            "hamming_threshold": "hamming图片相似度(在0到64之间)，一般为5的整数，数值越小，要求越高",
            "similarity_threshold": "similarity图片相似度(在0到1之间)，一般为0.9的浮点数，数值越大，要求越高",
            "tolerance": "表格单元格重复率(在0到1之间)，一般为0.9的浮点数，数值越大，要求越高",
            "min_chars": "文本连续最字符串值(大于等于1)，一般为50的整数，数值越小，谨慎度越高，相应地比对效率越低",
            "batch_size": "每次处理的任务数量(大于等于20)，默认为100的整数，数值越大，内存占用可能越高",
            "filter_chars": "您需要过滤掉的字符，如果是多个，用“；”分号隔开",
            }

        # hamming_threshold 设置
        label_hamming = tk.Label(settings_window, text="海明阈值 (hamming_threshold):")
        label_hamming.grid(row=0, column=0, padx=10, pady=5, sticky=tk.W)
        tooltip_hamming = tk.Label(settings_window, text="?", fg="blue", cursor="question_arrow")
        tooltip_hamming.grid(row=0, column=1, padx=(0, 10), pady=5, sticky=tk.W)  # 放置在label旁边
        entry_hamming = tk.Entry(settings_window, textvariable=self.hamming_threshold)
        entry_hamming.grid(row=0, column=2, padx=10, pady=5, sticky=tk.W)
        self.create_tooltip(tooltip_hamming, tooltips["hamming_threshold"])

        # similarity_threshold 设置
        label_similarity = tk.Label(settings_window, text="相似度阈值 (similarity_threshold):")
        label_similarity.grid(row=1, column=0, padx=10, pady=5, sticky=tk.W)
        tooltip_similarity = tk.Label(settings_window, text="?", fg="blue", cursor="question_arrow")
        tooltip_similarity.grid(row=1, column=1, padx=(0, 10), pady=5, sticky=tk.W)  # 放置在label旁边
        entry_similarity = tk.Entry(settings_window, textvariable=self.similarity)
        entry_similarity.grid(row=1, column=2, padx=10, pady=5, sticky=tk.W)
        self.create_tooltip(tooltip_similarity, tooltips["similarity_threshold"])

        # tolerance 设置
        label_tolerance = tk.Label(settings_window, text="容差 (tolerance):")
        label_tolerance.grid(row=2, column=0, padx=10, pady=5, sticky=tk.W)
        tooltip_tolerance = tk.Label(settings_window, text="?", fg="blue", cursor="question_arrow")
        tooltip_tolerance.grid(row=2, column=1, padx=(0, 10), pady=5, sticky=tk.W)  # 放置在label旁边
        entry_tolerance = tk.Entry(settings_window, textvariable=self.tolerance)
        entry_tolerance.grid(row=2, column=2, padx=10, pady=5, sticky=tk.W)
        self.create_tooltip(tooltip_tolerance, tooltips["tolerance"])

        # min_chars 设置
        label_min_chars = tk.Label(settings_window, text="最小字符数 (min_chars):")
        label_min_chars.grid(row=3, column=0, padx=10, pady=5, sticky=tk.W)
        tooltip_min_chars = tk.Label(settings_window, text="?", fg="blue", cursor="question_arrow")
        tooltip_min_chars.grid(row=3, column=1, padx=(0, 10), pady=5, sticky=tk.W)  # 放置在label旁边
        entry_min_chars = tk.Entry(settings_window, textvariable=self.min_chars)
        entry_min_chars.grid(row=3, column=2, padx=10, pady=5, sticky=tk.W)
        self.create_tooltip(tooltip_min_chars, tooltips["min_chars"])

        # batch_size 设置
        label_batch_size = tk.Label(settings_window, text="批次大小 (batch_size):")
        label_batch_size.grid(row=4, column=0, padx=10, pady=5, sticky=tk.W)
        tooltip_batch_size = tk.Label(settings_window, text="?", fg="blue", cursor="question_arrow")
        tooltip_batch_size.grid(row=4, column=1, padx=(0, 10), pady=5, sticky=tk.W)  # 放置在label旁边
        entry_batch_size = tk.Entry(settings_window, textvariable=self.batch_size)
        entry_batch_size.grid(row=4, column=2, padx=10, pady=5, sticky=tk.W)
        self.create_tooltip(tooltip_batch_size, tooltips["batch_size"])


        # filter_chars 设置
        label_filter_chars = tk.Label(settings_window, text="过滤字符 (filter_chars):")
        label_filter_chars.grid(row=5, column=0, padx=10, pady=5, sticky=tk.W)
        tooltip_filter_chars = tk.Label(settings_window, text="?", fg="blue", cursor="question_arrow")
        tooltip_filter_chars.grid(row=5, column=1, padx=(0, 10), pady=5, sticky=tk.W)  # 放置在label旁边
        entry_filter_chars = tk.Entry(settings_window, textvariable=self.filter_chars_var)
        entry_filter_chars.grid(row=5, column=2, padx=10, pady=5, sticky=tk.W)
        self.filter_chars_var.set(''.join(self.filter_chars) if self.filter_chars else '')  # 插入当前过滤字符
        self.create_tooltip(tooltip_filter_chars, tooltips["filter_chars"])


        # 添加移除所有标点符号的复选框
        check_remove_punctuation = Checkbutton(settings_window, text="过滤所有标点符号（一般不需要，除非你特想忽略掉所有标点符号进行文字比对！）", variable=self.remove_punctuation)
        check_remove_punctuation.grid(row=6, column=0, columnspan=3, padx=10, pady=5, sticky=tk.W)




        # 确认按钮
        tk.Button(settings_window, text="确认", command=lambda: self.update_settings(settings_window)).grid(row=7,
                                                                                                            column=0,
                                                                                                        columnspan=3,
                                                                                                            pady=10)

        # 设置窗口大小不可变
        settings_window.resizable(False, False)

    def update_settings(self, settings_window):
        # 更新 filter_chars 属性
        new_filter_chars = set()

        # 如果选择了移除所有标点符号，则添加所有标点符号到过滤字符集合
        if self.remove_punctuation.get():
            new_filter_chars.update(string.punctuation)  # 英文标点符号
            new_filter_chars.update("，。！？；：“”‘’（）【】《》、")  # 中文标点符号  01.15增加了顿号

        # 解析用户输入的过滤字符，并添加到集合中
        user_input_chars = self.filter_chars_var.get().strip()
        if user_input_chars:
            # 替换中文分号为英文分号，并分割字符串
            items = [item.strip() for item in user_input_chars.replace('；', ';').split(';') if item.strip()]

            for item in items:
                # 直接添加每个项中的所有字符到集合中
                for char in item:
                    new_filter_chars.add(char)

        self.filter_chars = new_filter_chars

        # 关闭设置窗口
        settings_window.destroy()


    def select_files(self):
        new_file_paths = filedialog.askopenfilenames(filetypes=[("Word Documents", "*.docx")])
        if new_file_paths:  # 检查是否有文件被选择
            self.file_paths.extend(new_file_paths)  # 将新选择的文件添加到列表中
            self.info_text.insert(tk.END, f"选择了文件: {', '.join(new_file_paths)}\n")
        else:
            self.info_text.insert(tk.END, "未选择任何文件\n")

    def clear_files(self):
        self.file_paths = []  # 清空文件列表
        self.info_text.delete('1.0', tk.END)  # 清空文本框内容
        self.info_text.insert(tk.END, "文件列表已清空\n")

    def select_output_path(self):
        new_output_path = filedialog.askdirectory()
        if new_output_path:  # 检查是否有路径被选择
            self.output_path = new_output_path  # 更新输出路径
            self.info_text.insert(tk.END, f"输出路径设置为: {self.output_path}\n")
        else:
            self.info_text.insert(tk.END, "未选择输出路径，保持之前的输出路径\n")

    def start_comparison(self):
        try:
            # 获取并验证用户输入的参数
            min_chars = self.min_chars.get()
            try:
                min_chars = int(min_chars)  # 尝试将输入转换为整数
                if min_chars < 1:
                    raise ValueError("min_chars必须大于等于1")
            except ValueError:
                raise ValueError("min_chars必须是大于等于1的整数")

            tolerance = self.tolerance.get()
            try:
                tolerance = float(tolerance)  # 尝试将输入转换为浮点数
                if not (0 <= tolerance <= 1):
                    raise ValueError("tolerance必须在0到1之间")
            except ValueError:
                raise ValueError("tolerance必须是有效的0到1之间浮点数")

            hamming_threshold = self.hamming_threshold.get()
            try:
                hamming_threshold = int(hamming_threshold)  # 尝试将输入转换为整数
                if not (0 <= hamming_threshold <= 64):
                    raise ValueError("hamming_threshold必须在0到64之间（包括0和64）。建议值为5")
            except ValueError:
                raise ValueError("hamming_threshold必须是在0到64之间的整数")


            similarity = self.similarity.get()
            try:
                similarity = float(similarity)  # 尝试将输入转换为浮点数
                if not (0 <= similarity <= 1):
                    raise ValueError("similarity_threshold必须在0到1之间")
            except ValueError:
                raise ValueError("similarity_threshold必须是在0到1之间浮点数")

            batch_size = self.batch_size.get()
            try:
                batch_size = int(batch_size)  # 尝试将输入转换为整数
                if batch_size < 20:
                    raise ValueError("batch_size必须大于等于20整数")
            except ValueError:
                raise ValueError("batch_size必须是大于等于20整数")

            # 使用已保存的过滤字符
            filter_chars = self.filter_chars

            if len(self.file_paths) < 2:
                messagebox.showwarning("警告", "请至少导入两个文件进行查重")
                return
            if not self.file_paths:
                messagebox.showwarning("警告", "请选择要比较的文件")
                return
            if not self.output_path:
                messagebox.showwarning("警告", "请选择输出路径")
                return

            self.progress_var.set(0)
            self.info_text.insert(tk.END, "开始比较.......\n请耐心等待......\n")
            self.info_text.see(tk.END)

            # 正确引用DocumentComparer类，并设置类级别的hamming_threshold
            setattr(DocumentComparer, 'hamming_threshold', hamming_threshold)
            # 正确引用DocumentComparer类
            self.comparer = DocumentComparer(min_chars=min_chars,tolerance=tolerance,hamming_threshold=hamming_threshold,similarity=similarity,batch_size = self.batch_size,file_paths=self.file_paths,
            output_path=self.output_path,filter_chars=filter_chars)
            self.progress_queue = Queue()
            # 启动线程进行文档比较，并传递similarity_threshold和其他参数
            self.thread = Thread(
                target=self.comparer.compare_documents,
                args=(self.progress_queue,),
                kwargs={
                    'similarity_threshold': similarity,
                    'min_chars': min_chars,
                    'tolerance': tolerance,
                    'hamming_threshold': hamming_threshold,
                    'filter_chars': filter_chars,  # 添加filter_chars参数
                    'batch_size': batch_size

                }
            )
            self.thread.start()
            self.check_progress()
        except ValueError as e:
            messagebox.showerror("错误", str(e))

    def check_progress(self):
        try:
            while True:
                progress = self.progress_queue.get(timeout=0.1)
                self.progress_var.set(progress)
                if progress == 100:
                    self.info_text.insert(tk.END, "比较完成\n")
                    self.info_text.see(tk.END)
                    break
        except Empty:
            self.root.after(100, self.check_progress)
        except KeyboardInterrupt:
            print("程序被用户中断")
            # 这里可以添加任何需要的清理代码
            self.root.quit()  # 安全退出Tkinter主循环
            self.root.destroy()



def check_for_updates(current_version, update_url, root=None):
    try:
        # 从服务器获取最新版本号
        response = requests.get(f"{update_url}/version（勿删）.txt", timeout=30)
        response.raise_for_status()  # 如果响应包含HTTP错误状态码，则抛出HTTPError
        latest_version = response.text.strip()

        #logging.info(f"服务器 'version（勿删）.txt' 内容: {latest_version}")



        # 比较版本号，假设版本号格式为 'YY.MM.DD'
        def versiontuple(v):
            return tuple(int(part) for part in v.split('.'))

        current_version_tuple = versiontuple(current_version)
        latest_version_tuple = versiontuple(latest_version)

        # 比较两个版本号
        if current_version_tuple < latest_version_tuple:
            logging.info(f"有新版本可用: {latest_version}")

            # 弹出消息框询问用户是否更新
            if root is None:
                root = tk.Tk()
                root.withdraw()  # 隐藏主窗口

            # 弹出消息框询问用户是否更新
            user_choice = messagebox.askyesno("更新提示", f"发现新版本 {latest_version}。您现在要更新吗？")

            if user_choice:
                logging.info("用户选择更新。")
                # 更新本地的 version.txt 文件
                current_script_dir = os.path.dirname(os.path.abspath(sys.argv[0]))
                version_txt_path = os.path.join(current_script_dir, "version（勿删）.txt")
                try:
                    with open(version_txt_path, 'w', encoding='utf-8') as f:
                        f.write(latest_version)
                    logging.info("成功更新 version（勿删）.txt。")
                except Exception as e:
                    logging.error(f"无法更新本地 version（勿删）.txt 文件: {e}")
                    return False, latest_version

                return True, latest_version
            else:
                logging.info("用户拒绝了更新。")
                return False, latest_version
        else:
            logging.info("没有可用的更新。")
            return False, None
    except requests.exceptions.RequestException as e:
        logging.info("无法检查更新，请检查您的网络连接（若您的网络连接无误，忽略此信息）。")
        return False, None

def download_and_replace(update_url_placeholder="[UPDATE_URL]", app_base_name="文档查重", doc_name="操作手册.docx", version=None):
    # 创建一个简单的 session 对象
    session = requests.Session()
    try:

        current_script_path = sys.argv[0]
        current_script_dir = os.path.dirname(os.path.abspath(current_script_path))

        doc_path = os.path.join(current_script_dir, doc_name)  # 在这里定义 doc_path

        # 下载操作手册.docx 文件
        try:
            doc_response = session.get(f"{update_url_placeholder}/{doc_name}", stream=True, timeout=None)
            doc_response.raise_for_status()
            with open(doc_path, 'wb') as f:
                for chunk in doc_response.iter_content(chunk_size=8192):
                    if chunk:
                        f.write(chunk)
            logging.info("成功下载操作手册.docx。")

            # 删除旧版本的操作手册.docx 文件
            old_doc_path = os.path.join(current_script_dir, doc_name)
            if os.path.exists(old_doc_path) and old_doc_path != doc_path:  # 确保不是同一文件
                os.remove(old_doc_path)
                logging.info("已删除旧的操作手册.docx 文件。")
            else:
                logging.info("未找到旧的操作手册.docx 文件或新旧文件路径相同。")
        except requests.exceptions.RequestException as e:
            logging.info("在下载操作手册.docx 文件时发生错误。")
            return False


        # 构建应用程序文件名
        default_app_name = f"{app_base_name}.exe"
        default_app_path = os.path.join(current_script_dir, default_app_name)

        # 分块下载新版本 exe 文件
        try:
            response = session.get(f"{update_url_placeholder}/{default_app_name}", stream=True, timeout=None)
            response.raise_for_status()
            with open(default_app_path, 'wb') as f:
                for chunk in response.iter_content(chunk_size=8192):
                    if chunk:
                        f.write(chunk)
            logging.info(f"成功下载新的应用程序文件 {default_app_name}。")
        except requests.exceptions.RequestException as e:
            logging.info("在下载应用程序文件时发生错误。")
            return False

        # 如果提供了版本号，则重命名文件以包含版本信息
        if version:
            new_app_name = f"{app_base_name}{version}.exe"
            new_app_path = os.path.join(current_script_dir, new_app_name)
            try:
                os.rename(default_app_path, new_app_path)
                logging.info(f"已将下载的文件重命名为 {new_app_name}。")
            except OSError as e:
                logging.error(f"无法重命名下载的文件: {e}")
                return False


        # 检查下载的 .exe 文件大小
        try:
            file_size = os.path.getsize(new_app_path)
            size_limit = 5 * 1024 * 1024  # 5 MB，如果下载下来的exe文件小于5MB，就停止更新
            if file_size < size_limit:
                logging.error(f"下载的文件大小 ({file_size / (1024*1024):.2f} MB) 小于5MB的阈值。更新已中止。")
                return False
            else:
                logging.info("下载的文件符合大小要求。")
        except OSError as e:
            logging.info("无法检查文件大小。")
            return False


        # 提示用户手动删除旧版本 exe 文件
        messagebox.showinfo("更新提示",
                            f"新版本的应用程序已下载完成 ({new_app_name})。\n请关闭当前应用程序，运行新版本无误后手动删除旧版本。")
        logging.info("已提示用户手动删除旧版本应用程序文件。")



        logging.info("更新已成功完成。")
        return True
    # except Exception as e:
    #     logging.info(f"Error downloading and replacing the application or version.txt: {e}")
    #     return False

    except requests.exceptions.RequestException as e:
        logging.info("检查更新失败。")
        return False
    finally:
        # 确保关闭session
        session.close()

def ensure_version_file_exists(current_version, script_dir):
    version_txt_path = os.path.join(script_dir, "version（勿删）.txt")
    if not os.path.exists(version_txt_path):
        with open(version_txt_path, 'w', encoding='utf-8') as f:
            f.write(current_version)
        logging.info(f"已使用版本{current_version}创建默认的 version（勿删）.txt")

def perform_update_check_and_install(current_version, update_url, root=None):
    has_update, latest_version = check_for_updates(current_version, update_url, root)
    if has_update:
        if download_and_replace(update_url, app_base_name="文档查重", doc_name="操作手册.docx",version=latest_version):
            logging.info("应用程序已更新。")

        else:
            logging.info("下载更新失败，不进行更新直接启动应用程序。")
    elif latest_version is None:
        logging.info("无法更新，直接启动应用程序。")




def main():
    #由于使用了多线程，所以打包exe时出现了多模块不能导入，需要在这里加这条语句，还要在开始就导入multiprocessing
    multiprocessing.freeze_support()


    # 获取当前脚本所在的目录
    script_dir = os.path.dirname(os.path.abspath(sys.argv[0]))


    # 确保 version.txt 文件存在
    current_version = "25.01.22"  # 替换为实际的当前版本号
    ensure_version_file_exists(current_version, script_dir)


    # 创建 Tkinter 根窗口
    root = tk.Tk()
    root.title("应用程序")

    # 执行更新检查
    update_url = "https://416391746.github.io/update-files"
    perform_update_check_and_install(current_version, update_url)





    # 创建 EmailVerifier 实例
    email_verifier = EmailVerifier()


    # 检查验证状态
    if not email_verifier.check_verification_status():
        # 如果未验证，提示用户进行验证
        email_verifier.ask_for_code_and_verify(root)
    else:
        # 如果已验证，启动  主程序
        same_word_app = DocumentApp(root)
    try:
        root.mainloop()
    except KeyboardInterrupt:
        print("程序被用户中断")
        # 确保事件循环停止，并且仅在root存在的情况下尝试销毁窗口
        if root is not None:
            try:
                root.quit()  # 停止事件循环
                root.destroy()  # 销毁窗口
            except tk.TclError:
                # 如果窗口已经被销毁，则忽略此错误
                pass
    except Exception as e:
        logging.error(f"发生意外错误: {e}", exc_info=True)
        if root is not None:
            try:
                root.quit()
                root.destroy()
            except tk.TclError:
                pass

if __name__ == "__main__":
    main()
